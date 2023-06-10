import torch
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
import time


class GenReport():
    def __init__(self, saved_path):
        super(GenReport, self).__init__()
        self.saved_path = saved_path
        self.document = Document()
        self.current_time = time.ctime()
        self.document.add_heading('DGFDBenchmark Tune Report', level=0) # level=0，文档标题
        self.document.add_paragraph('Report Date:'+ self.current_time)

    def _gen_figure(self, loss_acc_result):
        '''
        https://blog.csdn.net/ccc369639963/article/details/122980757
        '''
        fig1 = plt.figure()
        ax1 = fig1.add_axes([0,0,1,1])
        loss_keys=[]
        for key in loss_acc_result.keys():
            if key.__contains__('loss'):
                ax1.plot(loss_acc_result[key])
                loss_keys.append(key)
        # l11 = ax1.plot(loss_acc_result['loss_rr'],'y-')
        # l12 = ax1.plot(loss_acc_result['loss_cd'],'g-')
        # l13 = ax1.plot(loss_acc_result['loss_ca'],'m-')
        # l14 = ax1.plot(loss_acc_result['loss_cl'],'b-')
        # l15 = ax1.plot(loss_acc_result['loss_rr'] + loss_acc_result['loss_cd']+ loss_acc_result['loss_ca']+ loss_acc_result['loss_cl'],'r-')
        # ax1.legend(labels = ('loss_rr', 'loss_cd', 'loss_ca', 'loss_cl', 'total loss'))
        ax1.legend(labels = tuple(loss_keys))

        ax1.set_title('Loss Curves')
        ax1.set_xlabel('Steps')
        ax1.set_ylabel('Loss')
        ax1.grid(True)
        #https://matplotlib.org/stable/api/_as_gen/matplotlib.axes.Axes.legend.html
        fig2 = plt.figure()
        acces = loss_acc_result['acces']
        ax2 = fig2.add_axes([0,0,1,1])
        # l21 = ax2.plot(acces[:,0],'ys-')
        # l22 = ax2.plot(acces[:,1],'go--')
        l21 = ax2.plot(acces[:,0],'ys-')
        l22 = ax2.plot(np.mean(acces[:,1:],1),'go--')

        ax2.legend(labels = ('test_acc', 'train_acc'))
        ax2.set_title('Training and Test Accuracy')
        ax2.set_xlabel('Steps')
        ax2.set_ylabel('Accuracy')
        ax2.grid(True)

        return fig1, fig2

    def _write_figure(self, tmp_path, loss_acc_result):
        fig1, fig2 = self._gen_figure(loss_acc_result)
        fig1.savefig(tmp_path+'the_loss_fig_tmp.jpg', format='jpg',dpi=200, bbox_inches = 'tight')
        fig2.savefig(tmp_path+'the_acc_fig_tmp.jpg', format='jpg',dpi=200,  bbox_inches = 'tight')
        self.document.add_picture(tmp_path+'the_loss_fig_tmp.jpg')
        self.document.add_picture(tmp_path+'the_acc_fig_tmp.jpg')


    def _write_configs(self, configs):
        cfg = []
        for k, v in sorted(vars(configs).items()):
            cfg.append((k,v))

        # table = self.document.add_table(rows=len(cfg), cols=2)
        table = self.document.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells # the header of the table
        hdr_cells[0].text = 'Parameter Name'
        hdr_cells[1].text = 'Parameter Value'
        for para_name, para_value in cfg:
            row_cells = table.add_row().cells
            row_cells[0].text = para_name
            row_cells[1].text = str(para_value)

    def write_file(self, configs, test_item, loss_acc_result):
        tmp_path = self.saved_path
        self.document.add_heading('Test'+str(test_item), level=1)
        self.document.add_paragraph('Configs')
        self._write_configs(configs)
        self._write_figure(tmp_path, loss_acc_result)

    def save_file(self, currtime):
        self.document.save(self.saved_path+'//demo'+ currtime+'.docx')
